import docx
from docx.shared import Pt, RGBColor, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

def add_proper_algorithm_table(doc, title, lines):
    p_title = doc.add_paragraph()
    runner = p_title.add_run(title)
    runner.bold = True
    p_title.style = 'Caption'
    
    table = doc.add_table(rows=0, cols=1)
    table.style = 'Table Grid'
    
    for i, line in enumerate(lines, 1):
        row = table.add_row()
        cell = row.cells[0]
        if ":" in line and line.split(":")[0].isdigit():
            lnum, text = line.split(":", 1)
            p = cell.add_paragraph()
            run_lnum = p.add_run(f"{lnum:2s}: ")
            run_lnum.font.name = 'Courier New'
            run_lnum.font.size = Pt(9)
            run_lnum.bold = True
            
            run_text = p.add_run(text)
            run_text.font.name = 'Courier New'
            run_text.font.size = Pt(10)
        else:
            p = cell.add_paragraph()
            if line.startswith("Input:") or line.startswith("Output:") or line.startswith("Require:") or line.startswith("Ensure:"):
                kw, rest = line.split(":", 1)
                run_kw = p.add_run(f"{kw}:")
                run_kw.bold = True
                run_rest = p.add_run(rest)
            else:
                p.add_run(line)
                
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)
        if "    " in line:
            p.paragraph_format.left_indent = Inches(0.2)

doc_path = r'd:\web_attack_detector\Journal_Paper_Hybrid_Quantum_Web_Attack_Detector_Humanized.docx'
out_path = r'd:\web_attack_detector\Journal_Paper_Hybrid_Quantum_Web_Attack_Detector_Expanded.docx'

doc = docx.Document(doc_path)

paragraphs = list(doc.paragraphs)
for p in paragraphs:
    if p.text.startswith("Algorithm 1: Advanced Feature") or \
       p.text.startswith("Algorithm 2: Hybrid Quantum") or \
       p.text.startswith("Input: Raw HTTP Request") or \
       p.text.startswith("Output: 30-Dimensional") or \
       p.text.startswith("Input: Normalized Feature") or \
       p.text.startswith("Output: Definite Classification"):
        p._element.getparent().remove(p._element)
    elif "F.append(" in p.text or "SQL_Keywords <-" in p.text or "Anomaly_Flag <-" in p.text or "Weighted_Proba <-" in p.text:
       p._element.getparent().remove(p._element)

doc.add_heading('Formatted Algorithms', level=1)

add_proper_algorithm_table(doc, "Algorithm 1: Advanced Feature Extraction from Call Packets", [
    "Require: Raw HTTP Request Dictionary (URL, Headers, Body)",
    "Ensure: 30-Dimensional Feature Vector F",
    "1: Initialize empty array F",
    "2: F.append( calculate_string_length(URL) )",
    "3: F.append( compute_shannon_entropy(URL) )",
    "4: SQL_Keywords <- count(URL, ['SELECT', 'UNION', 'DROP'])",
    "5: XSS_Keywords <- count(URL, ['<script>', 'onerror=', 'alert'])",
    "6: F.extend( [SQL_Keywords, XSS_Keywords] )",
    "7: For each Header in Headers:",
    "8:     Determine structural anomaly markers",
    "9:     F.extend_accumulated_metrics()",
    "10: Standardize F using Min-Max Scaler",
    "11: Return F"
])

doc.add_paragraph("\n")

add_proper_algorithm_table(doc, "Algorithm 2: Hybrid Quantum-Classical Ensemble Escalation", [
    "Require: Normalized Feature Vector X",
    "Ensure: Classification Label (y_hat) and Confidence (C)",
    "1: Anomaly_Flag <- IsolationForest.Predict(X)",
    "2: Classical_Label, Classical_Conf <- XGBoost_Ensemble(X)",
    "3: IF Classical_Conf < 0.85 THEN",
    "4:     Quantum_Proba <- PennyLane_VQC.Evaluate(X)",
    "5:     Classical_Proba <- XGBoost_Ensemble.Proba(X)",
    "6:     Weighted_Proba <- (Classical_Conf * Classical_Proba) + (Quantum_Conf * Quantum_Proba)",
    "7:     Weighted_Proba <- Weighted_Proba / (Classical_Conf + Quantum_Conf)",
    "8:     y_hat <- ARGMAX(Weighted_Proba)",
    "9:     C <- MAX(Weighted_Proba)",
    "10: ELSE",
    "11:    y_hat <- Classical_Label",
    "12:    C <- Classical_Conf",
    "13: IF Anomaly_Flag == TRUE AND y_hat == 'Normal' THEN",
    "14:    y_hat <- 'Suspicious'",
    "15: Return y_hat, C"
])

doc.save(out_path)
print(f"Docx algorithms refactored and saved to {out_path}.")
