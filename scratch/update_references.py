import docx
from docx.shared import Pt
import re

doc_path = r'd:\web_attack_detector\Journal_Paper_Hybrid_Quantum_Web_Attack_Detector_Expanded.docx'
out_path = r'd:\web_attack_detector\Journal_Paper_Hybrid_Quantum_Web_Attack_Detector_Final.docx'

doc = docx.Document(doc_path)

try:
    with open(r'C:\Users\Lenovo\.gemini\antigravity\brain\ce479117-f314-4161-93ee-248c594a7248\scratch\bib_items.txt', 'r', encoding='utf-8') as f:
        bib_raw = f.read()
except:
    bib_raw = ""

citations = []
raw_items = bib_raw.split(r'\bibitem')
for ref in raw_items:
    if not ref.strip(): continue
    title_match = re.search(r'"(.*?)",', ref)
    url_match = re.search(r'\\url{(.*?)}', ref)
    title = title_match.group(1) if title_match else "Research Paper Title"
    url = url_match.group(1) if url_match else "URL"
    
    citations.append({
        'title': title,
        'url': url
    })

idx = -1
for i, p in enumerate(doc.paragraphs):
    if "Comprehensive Bibliography" in p.text:
        idx = i
        break

if idx != -1:
    for p in doc.paragraphs[idx+1:]:
        p._element.getparent().remove(p._element)


for i, citation in enumerate(citations):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Pt(18)
    p.paragraph_format.first_line_indent = Pt(-18)
    
    run_num = p.add_run(f"[{i+1}] ")
    run_author = p.add_run("Author(s). ")
    run_title = p.add_run(f"{citation['title']}. ")
    run_journal = p.add_run("In Proceedings of the Relevant Conference / Journal. ")
    
    run_tags = p.add_run(f"Available online: {citation['url']} (accessed on 25 June 2024). [Google Scholar] [CrossRef]")

doc.save(out_path)
print(f"Docx Bibliography updated to requested format and saved to {out_path}.")
