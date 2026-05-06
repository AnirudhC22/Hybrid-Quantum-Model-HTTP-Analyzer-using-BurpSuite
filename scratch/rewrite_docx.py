import docx
from docx.shared import Pt, Inches
import re
import os

doc_path = r'd:\web_attack_detector\Journal_Paper_Hybrid_Quantum_Web_Attack_Detector_Final.docx'
out_path = r'd:\web_attack_detector\Journal_Paper_Hybrid_Quantum_Web_Attack_Detector_Final_Polished.docx'

doc = docx.Document(doc_path)

# Insert the Executive Summary right after Abstract (or find first heading)
for i, p in enumerate(doc.paragraphs):
    if p.text.startswith("Abstract"):
        # We'll insert after the abstract paragraph itself
        idx = i + 2
        
        # We need to use insert_paragraph_before on the next paragraph
        if idx < len(doc.paragraphs):
            next_p = doc.paragraphs[idx]
            
            p1 = next_p.insert_paragraph_before("Executive Summary", style='Heading 1')
            p2 = next_p.insert_paragraph_before("Quantum ML is on the rise. Recent years have seen a surge of quantum-classical models for pattern classification and anomaly detection. For example, Havlíček et al. (2019) introduced a hybrid variational quantum classifier and quantum kernel that operate in a quantum-enhanced feature space, and Beer et al. (2020) defined a quantum perceptron and training algorithm that scales with network width...")
            p3 = next_p.insert_paragraph_before("Hybrid quantum-classical IDS are emerging. A variety of hybrid systems combine classical feature processing with small quantum subnets. For instance, QCNN-ID embeds a small quantum convolutional neural layer within a CNN for IoT traffic...")
            p4 = next_p.insert_paragraph_before("Selective routing & confidence-gating is nascent. The idea of 'selective intelligence routing' – using a cheap classical classifier first and routing only low-confidence cases to a heavy model – is relatively unexplored in QIDS...")
            p5 = next_p.insert_paragraph_before("Datasets and real-world practice. Many quantum IDS studies use legacy datasets (KDD-Cup99, NSL-KDD)...")
            p6 = next_p.insert_paragraph_before("Top Recommendations:")
            p6.bold = True
            next_p.insert_paragraph_before("1. Havlíček et al. (Nature 2019) – introduced a practical variational quantum classifier (VQC)", style='List Number')
            next_p.insert_paragraph_before("2. Kim & Madhavi (Sci. Rep. 2024) – presented an outlier-based quantum intrusion detector", style='List Number')
            next_p.insert_paragraph_before("3. Sakthivel et al. (Sci. Rep. 2026) – demonstrated confidence-gated IDS (SVM cascade)", style='List Number')
            next_p.insert_paragraph_before("4. Cirillo & Esposito (ICAART 2025) – described a hybrid Quantum GAN IDS", style='List Number')
            next_p.insert_paragraph_before("5. Chaudhary et al. (arXiv 2025) – a broad survey of federated and quantum NIDS", style='List Number')
        break

# Insert Architecture Images
for p in doc.paragraphs:
    if p.text.startswith("Proposed Methodology: In-Depth Architecture"):
        # Insert image
        run = p.add_run()
        run.add_break()
        try:
            # We will generate this image later
            run.add_picture(r'd:\web_attack_detector\scratch\server_arch.png', width=Inches(6.0))
        except:
            run.add_text("[Architecture Diagram Missing]")

# Replace the Bibliography
# The user wants exact Citations from their table
refs_text = [
    "Havlíček et al. Supervised learning with quantum-enhanced feature spaces. Nature 2019.",
    "Cong et al. Quantum convolutional neural networks. Nat. Phys. 2019.",
    "Beer et al. Training deep quantum neural networks. Nat. Comm. 2020.",
    "Biamonte et al. Quantum machine learning. Nature 2017.",
    "Kukliansky et al. QNN for network anomaly detection on IonQ hardware. T-QE 2024.",
    "Kim & Madhavi. QML-based IDS via quantum outlier detection. Sci. Rep. 2024.",
    "Sakthivel & Kumarasamy. Confidence-gated IDS cascade. Sci. Rep. 2026.",
    "Amara et al. QCNN-ID: Hybrid IoT intrusion detection. 2025.",
    "Chaudhary et al. Survey: federated & quantum ML for NIDS. arXiv 2025.",
    "Cirillo & Esposito. Quantum GAN for NSL-KDD anomaly detection. ICAART 2025.",
    "Abreu et al. QML-IDS for network attacks. 2024.",
    "Abreu et al. QuantumNetSec hybrid QML for cybersecurity. 2024.",
    "Bharathi et al. Hybrid QML for IoT IDS using PCA and VQC. EPJQT 2026.",
    "Senthil & Wong. Quantum Autoencoder (QAE) for cybersecurity anomalies. 2025.",
    "Wang et al. Detecting unseen malicious traffic with Hybrid QCNN. 2025.",
    "Cotrupi & Callahan. Quantum kernel SVM on IDS. 2024.",
    "Chen et al. Cyber threat classification via Hybrid QSVM/VQC on NSL-KDD. 2025.",
    "Nalayini et al. SDN intrusion detection via ATQ-IDS (Transformer + QIES). Sci Rep 2025."
]

idx = -1
for i, p in enumerate(doc.paragraphs):
    if "Comprehensive Bibliography" in p.text:
        idx = i
        break

if idx != -1:
    for p in doc.paragraphs[idx+1:]:
        p._element.getparent().remove(p._element)

    for i, citation in enumerate(refs_text):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Pt(18)
        p.paragraph_format.first_line_indent = Pt(-18)
        
        run_num = p.add_run(f"[{i+1}] ")
        
        # Split out author and title if possible
        parts = citation.split(".", 1)
        if len(parts) > 1:
            run_author = p.add_run(parts[0] + ".")
            run_title = p.add_run(parts[1])
            run_title.italic = True
        else:
            p.add_run(citation)
            
        p.add_run(" [Google Scholar] [CrossRef]")

doc.save(out_path)
print(f"Docx with Executive Summary, Diagram hooks, and proper Citations saved to {out_path}.")
